#!/usr/bin/env python3
"""
–°–û–ó–î–ê–¢–ï–õ–¨ DOCX –î–û–ö–£–ú–ï–ù–¢–û–í
"""

import os
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =============================================================================
# –§–£–ù–ö–¶–ò–ò –î–õ–Ø –í–´–í–û–î–ê ASCII –ê–†–¢–ê
# =============================================================================

def print_art():
    """–í—ã–≤–æ–¥–∏—Ç –≤—Å—Ç—Ä–æ–µ–Ω–Ω—ã–π ASCII –∞—Ä—Ç"""
    art = """
‚ïî‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïó
‚ïë       –î–û–ö–£–ú–ï–ù–¢–û–ì–ï–ù–ï–†–ê–¢–û–† 3000       ‚ïë
‚ïë           –≤–µ—Ä—Å–∏—è 1.0.0              ‚ïë
‚ïö‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïù
    """
    print(art)
    return True

def clear_screen():
    """–û—á–∏—Å—Ç–∫–∞ —ç–∫—Ä–∞–Ω–∞"""
    os.system('cls' if os.name == 'nt' else 'clear')

# =============================================================================
# –ö–õ–ê–°–° –î–õ–Ø –°–û–ó–î–ê–ù–ò–Ø DOCX –î–û–ö–£–ú–ï–ù–¢–û–í
# =============================================================================

class DocxCreator:
    def __init__(self):
        self.document = Document()
        self.setup_fixed_styles()
        
    def setup_fixed_styles(self):
        """–ù–∞—Å—Ç—Ä–æ–π–∫–∞ —Ñ–∏–∫—Å–∏—Ä–æ–≤–∞–Ω–Ω—ã—Ö —Å—Ç–∏–ª–µ–π –¥–æ–∫—É–º–µ–Ω—Ç–∞"""
        # –ö–æ–Ω—Å—Ç–∞–Ω—Ç—ã
        self.DEFAULT_FONT = 'Times New Roman'
        self.DEFAULT_COLOR = RGBColor(0, 0, 0)  # –ß–µ—Ä–Ω—ã–π
        self.LINE_SPACING = 1.15
        
        # –ù–∞—Å—Ç—Ä–æ–π–∫–∞ —Å—Ç—Ä–∞–Ω–∏—Ü—ã
        section = self.document.sections[0]
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21.0)
        section.left_margin = Cm(3.0)    # –õ–µ–≤—ã–π –æ—Ç—Å—Ç—É–ø 3 —Å–º
        section.right_margin = Cm(2.0)   # –ü—Ä–∞–≤—ã–π –æ—Ç—Å—Ç—É–ø 2 —Å–º
        section.top_margin = Cm(2.0)     # –í–µ—Ä—Ö–Ω–∏–π –æ—Ç—Å—Ç—É–ø 2 —Å–º
        section.bottom_margin = Cm(2.0)  # –ù–∏–∂–Ω–∏–π –æ—Ç—Å—Ç—É–ø 2 —Å–º
        
        # –ù–∞—Å—Ç—Ä–æ–π–∫–∞ —Å—Ç–∏–ª—è Normal
        style = self.document.styles['Normal']
        style.font.name = self.DEFAULT_FONT
        style.font.size = Pt(12)
        style.font.color.rgb = self.DEFAULT_COLOR
        style.paragraph_format.line_spacing = self.LINE_SPACING
        style.paragraph_format.space_before = Cm(0)
        style.paragraph_format.space_after = Cm(0.5)
        style.paragraph_format.first_line_indent = Cm(1.25)  # –ö—Ä–∞—Å–Ω–∞—è —Å—Ç—Ä–æ–∫–∞
    
    def create_document(self, text):
        """–°–æ–∑–¥–∞—Ç—å –¥–æ–∫—É–º–µ–Ω—Ç –∏–∑ —Ç–µ–∫—Å—Ç–∞"""
        # –†–∞–∑–¥–µ–ª—è–µ–º –Ω–∞ —Å—Ç—Ä–æ–∫–∏
        lines = text.strip().split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º —Ä–∞–∑—Ä—ã–≤ —Å—Ç—Ä–∞–Ω–∏—Ü—ã
            if line == '[PAGE_BREAK]':
                self.document.add_page_break()
                continue
            
            # –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º –∑–∞–≥–æ–ª–æ–≤–æ–∫ –¥–æ–∫—É–º–µ–Ω—Ç–∞
            if line.startswith('<title>') and line.endswith('</title>'):
                self.add_title(line)
                continue
            
            # –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º —Å—Ç—Ä–æ–∫—É —Å —Ç–µ–≥–∞–º–∏
            self.process_line(line)
    
    def process_line(self, line):
        """–û–±—Ä–∞–±–æ—Ç–∞—Ç—å –æ–¥–Ω—É —Å—Ç—Ä–æ–∫—É —Å —Ç–µ–≥–∞–º–∏"""
        line = line.strip()
        
        # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –µ—Å—Ç—å –ª–∏ –∑–∞–≥–æ–ª–æ–≤–æ–∫ h1-h4
        heading_match = re.search(r'<(h[1-4])>(.*?)</\1>', line)
        if heading_match:
            heading_type = heading_match.group(1)  # h1, h2, h3, h4
            heading_content = heading_match.group(2).strip()
            level = int(heading_type[1])  # –ò–∑–≤–ª–µ–∫–∞–µ–º —Ü–∏—Ñ—Ä—É
            
            # –ò–∑–≤–ª–µ–∫–∞–µ–º –≤—ã—Ä–∞–≤–Ω–∏–≤–∞–Ω–∏–µ
            alignment = self.get_alignment(line)
            
            self.add_heading(heading_content, level, alignment)
        else:
            # –≠—Ç–æ –æ–±—ã—á–Ω—ã–π –∞–±–∑–∞—Ü
            alignment = self.get_alignment(line)
            clean_line = self.remove_alignment_tags(line)
            
            if clean_line:
                self.add_paragraph(clean_line, alignment)
    
    def get_alignment(self, line):
        """–û–ø—Ä–µ–¥–µ–ª–∏—Ç—å –≤—ã—Ä–∞–≤–Ω–∏–≤–∞–Ω–∏–µ –∏–∑ —Ç–µ–≥–æ–≤"""
        if '<c>' in line and '</c>' in line:
            return 'center'
        elif '<l>' in line and '</l>' in line:
            return 'left'
        elif '<p>' in line and '</p>' in line:
            return 'right'
        elif '<j>' in line and '</j>' in line:
            return 'justify'
        else:
            return 'justify'  # –ü–æ —É–º–æ–ª—á–∞–Ω–∏—é
    
    def remove_alignment_tags(self, line):
        """–£–±—Ä–∞—Ç—å —Ç–µ–≥–∏ –≤—ã—Ä–∞–≤–Ω–∏–≤–∞–Ω–∏—è –∏–∑ —Å—Ç—Ä–æ–∫–∏"""
        line = re.sub(r'</?(c|l|p|j)>', '', line)
        return line.strip()
    
    def add_title(self, text):
        """–î–æ–±–∞–≤–∏—Ç—å –∑–∞–≥–æ–ª–æ–≤–æ–∫ –¥–æ–∫—É–º–µ–Ω—Ç–∞"""
        title_match = re.search(r'<title>(.*?)</title>', text)
        if title_match:
            title_text = title_match.group(1).strip()
            
            paragraph = self.document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Cm(1.0)
            
            run = paragraph.add_run(title_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(20)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            self.document.add_paragraph()  # –ü—É—Å—Ç–∞—è —Å—Ç—Ä–æ–∫–∞
    
    def add_heading(self, text, level, alignment='left'):
        """–î–æ–±–∞–≤–∏—Ç—å –∑–∞–≥–æ–ª–æ–≤–æ–∫"""
        params = {
            1: {'size': 14, 'bold': False},
            2: {'size': 14, 'bold': True},
            3: {'size': 16, 'bold': True},
            4: {'size': 18, 'bold': True}
        }.get(level, {'size': 14, 'bold': False})
        
        paragraph = self.document.add_paragraph()
        
        # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –≤—ã—Ä–∞–≤–Ω–∏–≤–∞–Ω–∏–µ
        if alignment == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # –ù–∞—Å—Ç—Ä–∞–∏–≤–∞–µ–º –æ—Ç—Å—Ç—É–ø—ã
        paragraph.paragraph_format.space_before = Cm(0.3)
        paragraph.paragraph_format.space_after = Cm(0.3)
        
        # –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º inline-—Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ
        self.process_inline_formatting(paragraph, text, params['size'], params['bold'])
    
    def add_paragraph(self, text, alignment='justify'):
        """–î–æ–±–∞–≤–∏—Ç—å –æ–±—ã—á–Ω—ã–π –∞–±–∑–∞—Ü"""
        paragraph = self.document.add_paragraph()
        
        # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –≤—ã—Ä–∞–≤–Ω–∏–≤–∞–Ω–∏–µ
        if alignment == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # –ù–∞—Å—Ç—Ä–∞–∏–≤–∞–µ–º –æ—Ç—Å—Ç—É–ø—ã
        paragraph.paragraph_format.first_line_indent = Cm(1.25)  # –ö—Ä–∞—Å–Ω–∞—è —Å—Ç—Ä–æ–∫–∞
        paragraph.paragraph_format.space_after = Cm(0.5)
        
        # –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º inline-—Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ
        self.process_inline_formatting(paragraph, text, 12, False)
    
    def process_inline_formatting(self, paragraph, text, size, bold):
        """–û–±—Ä–∞–±–æ—Ç–∞—Ç—å inline-—Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ"""
        parts = re.split(r'(<[^>]+>)', text)
        
        current_bold = bold
        current_italic = False
        current_underline = False
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('<') and part.endswith('>'):
                tag = part[1:-1].lower()
                if tag == 'b':
                    current_bold = True
                elif tag == '/b':
                    current_bold = False
                elif tag == 'i':
                    current_italic = True
                elif tag == '/i':
                    current_italic = False
                elif tag == 'z':
                    current_underline = True
                elif tag == '/z':
                    current_underline = False
            else:
                if part.strip():
                    run = paragraph.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(size)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.bold = current_bold
                    run.italic = current_italic
                    run.underline = current_underline
    
    def save(self, filename):
        """–°–æ—Ö—Ä–∞–Ω–∏—Ç—å –¥–æ–∫—É–º–µ–Ω—Ç"""
        if not filename.endswith('.docx'):
            filename += '.docx'
        
        self.document.save(filename)
        return filename

# =============================================================================
# –ì–õ–ê–í–ù–ê–Ø –§–£–ù–ö–¶–ò–Ø –ü–†–û–ì–†–ê–ú–ú–´
# =============================================================================

def main():
    """–ì–ª–∞–≤–Ω–∞—è —Ñ—É–Ω–∫—Ü–∏—è –ø—Ä–æ–≥—Ä–∞–º–º—ã"""
    # 1. –ß–∏—Å—Ç–∏–º —Ç–µ—Ä–º–∏–Ω–∞–ª
    clear_screen()
    
    # 2. –í—ã–≤–æ–¥–∏–º –∞—Ä—Ç
    print_art()
    print()  # –ü—É—Å—Ç–∞—è —Å—Ç—Ä–æ–∫–∞
    
    # 3. –ó–∞–ø—Ä–∞—à–∏–≤–∞–µ–º —Ç–µ–∫—Å—Ç –¥–ª—è —Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏—è
    print("–í–≤–µ–¥–∏—Ç–µ —Ç–µ–∫—Å—Ç –¥–ª—è —Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏—è:")
    print("(–î–ª—è –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è –≤–≤–æ–¥–∞ –Ω–∞–∂–º–∏—Ç–µ Ctrl+D)")
    print("-" * 40)
    
    lines = []
    try:
        while True:
            line = input()
            lines.append(line)
    except EOFError:
        # Ctrl+D –±—ã–ª –Ω–∞–∂–∞—Ç
        pass
    except KeyboardInterrupt:
        # Ctrl+C –±—ã–ª –Ω–∞–∂–∞—Ç
        print("\n\n‚ùå –í–≤–æ–¥ –ø—Ä–µ—Ä–≤–∞–Ω –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª–µ–º!")
        return
    
    text = '\n'.join(lines)
    
    if not text.strip():
        print("\n‚ùå –¢–µ–∫—Å—Ç –Ω–µ –≤–≤–µ–¥–µ–Ω!")
        return
    
    # 4. –ó–∞–ø—Ä–∞—à–∏–≤–∞–µ–º –Ω–∞–∑–≤–∞–Ω–∏–µ —Ñ–∞–π–ª–∞
    print("\n–í–≤–µ–¥–∏—Ç–µ –Ω–∞–∑–≤–∞–Ω–∏–µ —Ñ–∞–π–ª–∞ (–±–µ–∑ .docx):")
    filename = input("> ").strip()
    
    if not filename:
        import datetime
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"–¥–æ–∫—É–º–µ–Ω—Ç_{timestamp}"
    
    # 5. –°–æ–∑–¥–∞–µ–º –¥–æ–∫—É–º–µ–Ω—Ç
    print("\nüîÑ –°–æ–∑–¥–∞–Ω–∏–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞...")
    
    creator = DocxCreator()
    creator.create_document(text)
    saved_file = creator.save(filename)
    
    # 6. –í—ã–≤–æ–¥–∏–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç
    print(f"\n‚úÖ –§–∞–π–ª —Å–æ–∑–¥–∞–Ω: {saved_file}")
    print("–ü—Ä–æ–≥—Ä–∞–º–º–∞ –∑–∞–≤–µ—Ä—à–µ–Ω–∞.")

# =============================================================================
# –ó–ê–ü–£–°–ö –ü–†–û–ì–†–ê–ú–ú–´
# =============================================================================

if __name__ == "__main__":
    main()
